# import whisperx
import gc
import json
import os
import shutil

import gradio as gr
import whisper

from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain.vectorstores import OpenSearchVectorSearch, ElasticsearchStore
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.chat_models import ChatOpenAI, ChatOllama
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import FlashrankRerank
from langchain import PromptTemplate, LLMChain
from langchain.llms import Ollama
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from sentence_transformers import CrossEncoder
from docx import Document


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    return abstract_summary


def abstract_summary_extraction(transcription):
    llm = ChatOllama(streaming=True, model="qwen2:72b",
                     callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]))
    template = """
    [transcription] start:
    {transcription}
    [transcription] end
    ------------------
    You are an intelligent assistant skilled in summarizing and drafting meeting minutes. referring [transcription] as a meeting record. 
    Please generate a detailed and concise set of meeting minutes based on this [transcription]. 
    Ensure the minutes include the following sections and with proper format of meeting record:
    1. Meeting Overview: Include the date, time, location, attendees, and the chairperson of the meeting.
    2. Agenda Items: List the main topics discussed during the meeting.
    3. Discussion Details: Summarize the discussion points and key insights for each agenda item.
    4. Decisions and Action Items: List the decisions made and specific action items to be executed, including the responsible person and the deadline.
    5. Other Matters: Record any additional points of note.
    """
    QA_CHAIN_PROMPT = PromptTemplate(
        input_variables=["transcription"],
        template=template,
    )
    prompt = QA_CHAIN_PROMPT.format(transcription=transcription)
    # res = llm.stream(prompt)
    partial_message = ''
    for response in llm.stream(prompt):
        partial_message += response.content
    return partial_message


def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def get_creation_time(path):
    # Return the creation time of the file/folder
    return os.path.getctime(path)


def delete_old_tmp_folders():
    current_directory = os.getcwd()

    # List all directories in current directory
    tmp_folders = [os.path.join(current_directory, item) for item in os.listdir(current_directory)
                   if os.path.isdir(os.path.join(current_directory, item)) and item.startswith('tmp')]

    # Sort folders by creation time (newest first)
    tmp_folders.sort(key=get_creation_time, reverse=True)

    # Keep the 3 most recent folders and delete the rest
    for folder in tmp_folders[3:]:
        try:
            shutil.rmtree(folder)
            print(f"Deleted folder: {folder}")
        except Exception as e:
            print(f"Failed to delete {folder}: {e}")

def get_transcription(file):
    audio_file = file
    # model = whisper.load_model("small", device='cuda')
    # model = whisper.load_model("small", device='cuda')
    model = whisper.load_model("medium", device='cuda')

    result = model.transcribe(audio_file)
    print(result["text"])
    return meeting_minutes(result["text"])


# def get_transcription_x(file):
#     device = "cuda"
#     audio_file = file
#     batch_size = 16  # reduce if low on GPU mem
#     compute_type = "float16"  # change to "int8" if low on GPU mem (may reduce accuracy)
#
#     # 1. Transcribe with original whisper (batched)
#     model = whisperx.load_model("large-v2", device, compute_type=compute_type)
#
#     # save model to local path (optional)
#     # model_dir = "/path/"
#     # model = whisperx.load_model("large-v2", device, compute_type=compute_type, download_root=model_dir)
#
#     audio = whisperx.load_audio(audio_file)
#     result = model.transcribe(audio, batch_size=batch_size)
#     print(result)  # before alignment
#     print(len(result["segments"]))
#     # delete model if low on GPU resources
#     # import gc; gc.collect(); torch.cuda.empty_cache(); del model
#
#     # 2. Align whisper output
#     # model_a, metadata = whisperx.load_align_model(language_code=result["language"], device=device)
#     # result = whisperx.align(result["segments"], model_a, metadata, audio, device, return_char_alignments=False)
#     #
#     # print(result["segments"])  # after alignment
#     return meeting_minutes(result["segments"])
