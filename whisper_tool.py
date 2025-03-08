# import whisperx
import gc
import json
import os
import shutil

import gradio as gr
import whisper

from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain.vectorstores import OpenSearchVectorSearch, ElasticsearchStore
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.chat_models import ChatOpenAI, ChatOllama
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import FlashrankRerank
from langchain import PromptTemplate, LLMChain
from langchain.llms import Ollama
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from sentence_transformers import CrossEncoder
from docx import Document

llm = ChatOllama(model="qwen2.5:72b",
                 callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]))
def meeting_minutes(transcription):
    abstract_summary=""
    if len(transcription.split()) > 6000:
        print("________________________________________split transcription____________________________")
        transcription_list = split_transcript(transcription)
        minutes_list = []
        for t in transcription_list:
            minutes_list.append(abstract_summary_extraction(t))
        abstract_summary = merge_minutes(minutes_list)
    else:
        abstract_summary = abstract_summary_extraction(transcription)

    return abstract_summary



def split_transcript(transcript, chunk_size=3000, overlap=500):
    """
    将会议记录分成多个部分，每部分约 chunk_size 个字符，有 overlap 个字符的重叠
    """
    words = transcript.split()
    chunks = []
    start = 0

    while start < len(words):
        chunk = ' '.join(words[start:start + chunk_size])
        if len(chunk) < chunk_size and start + chunk_size < len(words):
            # 如果当前块小于目标大小且还有更多单词，则添加更多单词
            chunk = ' '.join(words[start:start + chunk_size + 1000])  # 添加额外单词以达到目标大小
        chunks.append(chunk)
        start += chunk_size - overlap

    return chunks

def abstract_summary_extraction(transcription):
    # llm = ChatOllama(streaming=True, model="qwen2:72b",
    #                  callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]))
    template = """
    [transcription] start:
    {transcription}
    [transcription] end
    ------------------
    You are an intelligent assistant skilled in summarizing and drafting user-friendly meeting minutes based on [transcription] as the meeting conversation record.
    
    Please write a clear, concise, and easy-to-read meeting summary, ensuring it flows naturally and feels conversational.
    
    The meeting minutes should include the following sections, but feel free to adapt the tone to make the text smooth and professional:
    
    1. **Meeting Overview**: Provide the date, time, location, attendees, and the chairperson.
    2. **Agenda Items**: List the main topics covered.
    3. **Discussion Summary**: Briefly summarize the key points from the discussions, focusing on insights and outcomes for each agenda item.
    4. **Decisions and Action Items**: Highlight the main decisions made and any action points, clearly stating who is responsible for each task and the expected deadlines.
    5. **Additional Notes**: Mention any other important points or follow-up items that came up during the meeting.
    
    Keep the tone professional but friendly, and ensure the minutes are easy to read and understand.

    """
    # template = """
    # [transcription] start:
    #
    # {transcription}
    #
    # [transcription] end
    #
    # You are an intelligent assistant skilled in summarizing and drafting English version meeting minutes. Refer to the [transcription] as the meeting conversation record. Please very carefully analyze [transcription], don't skip any topic, then generate detailed, concise, complete, and comprehensive meeting minutes based on this [transcription].
    #
    # Pay special attention to the every part of the transcription to accurately capture:
    # 1. The meeting's context and purpose
    # 2. Any introductory remarks or announcements
    # 3. The order and structure of the agenda
    #
    # Ensure the minutes include the following sections, formatted for proper, easy understanding, and readability:
    # Ensure clarity, completeness, and accuracy in the summary to reflect the true essence of the meeting.
    # Use bullet points or numbered lists where appropriate to enhance readability.
    # Emphasize important points and ensure there is no ambiguity in the recorded minutes.
    # Ensure the entire meeting content is covered, including both the front and later parts.
    #
    # 1. Meeting Overview:
    # - **Date:** (Extract from the transcription or infer if possible)
    # - **Time:** (Extract start and end time if mentioned)
    # - **Location:** (Identify if it's a physical location or virtual meeting)
    # - **Meeting Type:** (e.g., Weekly Status Update, Project Kickoff, etc.)
    # - **Attendees:** (List all participants mentioned, including late arrivals or early departures)
    # - **Chairperson/Facilitator:** (Identify who is leading the meeting)
    # - **Objective:** (Briefly state the main purpose of the meeting)
    #
    # 2. Opening Remarks:
    # - **Agenda Review:** (Note if the agenda was presented or reviewed at the start)
    # - **Previous Action Items:** (Mention any follow-ups from previous meetings discussed)
    #
    # 3. Agenda Items:
    # - **List of Main Topics:** (Provide a clear, ordered list of all discussed topics)
    #
    # 4. Discussion Details:
    # - **Summarize Key Points:** For each agenda item, provide a detailed summary of the discussion, including:
    #   - Key insights and perspectives shared
    #   - Important arguments or points made by specific attendees
    #   - Any relevant data or reports mentioned
    #
    # 5. Decisions and Action Items:
    # - **Decisions Made:** Clearly list all decisions made during the meeting, specifying:
    #   - The decision
    #   - The person(s) who made the decision
    #   - Any relevant context or reasoning behind the decision
    # - **Action Items:** Detail all action items to be executed, including:
    #   - A clear description of the task
    #   - The responsible person(s) or team
    #   - The deadline for completion
    #   - Any resources or support required
    #
    # 6. Other Matters:
    # - **Additional Points of Note:** Record any additional discussions or points of interest not covered under agenda items, including:
    #   - Announcements
    #   - Future meeting dates and times
    #   - Any follow-up required
    #   - Other relevant information or notes
    # """


    QA_CHAIN_PROMPT = PromptTemplate(
        input_variables=["transcription"],
        template=template,
    )
    prompt = QA_CHAIN_PROMPT.format(transcription=transcription)
    # res = llm.stream(prompt)
    partial_message = ''
    for response in llm.stream(prompt):
        partial_message += response.content
    return partial_message

def merge_minutes(minutes_list):
    """
    合并多个部分的会议纪要
    """
    merged = "Combined Meeting Minutes:\n\n"
    for i, minutes in enumerate(minutes_list, 1):
        merged += f"Part {i}:\n{minutes}\n\n-------------------------\n\n"

    merged += """
    Please review and consolidate the above parts of the meeting minutes, those meeting minutes belongs to same meeting, you task is merge them to one meeting minute.
    removing any duplications and ensuring a coherent flow of information. 
    Pay special attention to:
    1. Consistent formatting throughout the document
    2. Removing repeated information
    3. Ensuring all unique points from all parts are included
    4. Maintaining a logical order of topics and discussions
    5. Creating a unified document that reads as if it was created from a single transcript
    """
    prompt = merged


    # res = llm.stream(prompt)
    partial_message = ''
    for response in llm.stream(prompt):
        partial_message += response.content
    return partial_message

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def get_creation_time(path):
    # Return the creation time of the file/folder
    return os.path.getctime(path)


def delete_old_tmp_folders():
    current_directory = os.getcwd()

    # List all directories in current directory
    tmp_folders = [os.path.join(current_directory, item) for item in os.listdir(current_directory)
                   if os.path.isdir(os.path.join(current_directory, item)) and item.startswith('tmp')]

    # Sort folders by creation time (newest first)
    tmp_folders.sort(key=get_creation_time, reverse=True)

    # Keep the 3 most recent folders and delete the rest
    for folder in tmp_folders[3:]:
        try:
            shutil.rmtree(folder)
            print(f"Deleted folder: {folder}")
        except Exception as e:
            print(f"Failed to delete {folder}: {e}")

def get_transcription(file,model):
    audio_file = file
    # model = whisper.load_model("small", device='cuda')
    # model = whisper.load_model("small", device='cuda')

    result = model.transcribe(audio_file)
    print(result["text"])
    return meeting_minutes(result["text"])


# def get_transcription_x(file):
#     device = "cuda"
#     audio_file = file
#     batch_size = 16  # reduce if low on GPU mem
#     compute_type = "float16"  # change to "int8" if low on GPU mem (may reduce accuracy)
#
#     # 1. Transcribe with original whisper (batched)
#     model = whisperx.load_model("large-v2", device, compute_type=compute_type)
#
#     # save model to local path (optional)
#     # model_dir = "/path/"
#     # model = whisperx.load_model("large-v2", device, compute_type=compute_type, download_root=model_dir)
#
#     audio = whisperx.load_audio(audio_file)
#     result = model.transcribe(audio, batch_size=batch_size)
#     print(result)  # before alignment
#     print(len(result["segments"]))
#     # delete model if low on GPU resources
#     # import gc; gc.collect(); torch.cuda.empty_cache(); del model
#
#     # 2. Align whisper output
#     # model_a, metadata = whisperx.load_align_model(language_code=result["language"], device=device)
#     # result = whisperx.align(result["segments"], model_a, metadata, audio, device, return_char_alignments=False)
#     #
#     # print(result["segments"])  # after alignment
#     return meeting_minutes(result["segments"])
